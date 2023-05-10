import nltk
from nltk.tree import Tree
from nltk.corpus import wordnet
from nltk.draw import TreeWidget
from nltk.draw.tree import TreeView
from nltk.draw.util import CanvasFrame
import docx
import json
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
import tkinter.ttk as ttk
import string
from tkinter import Tk, Canvas, Text, Button, PhotoImage, Label, ttk, Scrollbar, filedialog, Toplevel
from PIL import ImageTk, Image
import PyPDF2


class Main():
    def __init__(self):
        self.init_main()
        self.new_doc = []
        self.result = {}
        self.curr_result = {}

    def save_button_click(self):
        self.save_list_to_file()

    def save_list_to_file(self):
        file_path = filedialog.asksaveasfilename(defaultextension='.json')
        with open(file_path, 'w') as f:
            json.dump(self.curr_result, f)

    def upload_list_from_file(self):
        file_path = filedialog.askopenfilename(filetypes=[('JSON Files', '*.json')])
        with open(file_path, 'r') as f:
            self.result = json.load(f)
            self.curr_result = self.result

    def upload_button_click(self):
        self.upload_list_from_file()
        print(self.result)
        self.show_curr_dict()

    def show_curr_dict(self):
        self.list_box.delete(0, END)

        k = 0
        for token in self.curr_result.keys():
            self.list_box.insert(k, '---слово---:' + token)
            self.list_box.insert(k+1, 'синонимы: ' + self.curr_result[token]['synonyms'])
            self.list_box.insert(k+2, 'антонимы: ' + self.curr_result[token]['antonyms'])
            self.list_box.insert(k+3, 'гипонимы: ' + self.curr_result[token]['hyponyms'])
            self.list_box.insert(k+4, 'гиперонимы: ' + self.curr_result[token]['hyperonyms'])
            k += 5

    def show_full_dict(self):
        self.list_box.delete(0, END)

        k = 0
        for token in self.result.keys():
            self.list_box.insert(k, '---слово---: ' + token)
            self.list_box.insert(k + 1, 'синонимы: ' + self.result[token]['synonyms'])
            self.list_box.insert(k + 2, 'антонимы: ' + self.result[token]['antonyms'])
            self.list_box.insert(k + 3, 'гипонимы: ' + self.result[token]['hyponyms'])
            self.list_box.insert(k + 4, 'гиперонимы: ' + self.result[token]['hyperonyms'])
            k += 5

        self.curr_result = self.result

    def back_button_click(self):
        self.show_full_dict()

    def update_button_click(self):
        child_update_empl(self.window, self, str(self.result))

    def document_button_click(self):
        child_document_empl(self.window, self, str(self.result))


    def addTextFromFile(self):
        #file_name = filedialog.askopenfilename(filetypes=[('MicrosoftWord' ,'*.docx ')])
        file_name = filedialog.askopenfilename(filetypes=[('Microsoft Edge PDF', '*.pdf')])
        if file_name != '':
            with open(file_name, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                text = '\n'
                for i in range(num_pages):
                    page = reader.pages[i]
                    text += page.extract_text()

                lines = text.splitlines()
                text = ' '.join(lines)
                self.inputText.delete(1.0, END)
                self.inputText.insert(1.0, text)
    
    

    def viewWindow(self):
        self.list_box.delete(0, END)
        text = self.inputText.get(1.0, END)
        text = text.replace('\n', '')
        print(text)

        self.result = {}
        self.curr_result = {}

        if text != '':
            tokens = text.split()
            tokens = [token.strip(string.punctuation) for token in tokens if token.strip(string.punctuation)]
            print(tokens)

            for token in tokens:
                self.result[token] = {}
                self.result[token]['synonyms'] = ''
                self.result[token]['antonyms'] = ''
                self.result[token]['hyponyms'] = ''
                self.result[token]['hyperonyms'] = ''
                print(token)
                syn = wordnet.synsets(token)
                print(syn)
                if len(syn) != 0:
                    for l in syn[0].lemmas():
                        self.result[token]['synonyms'] += l.name() + ', '
                        print(l.name)
                        if l.antonyms():
                            self.result[token]['antonyms'] += l.antonyms()[0].name() + ', '

                    for i in syn[0].hyponyms():
                        self.result[token]['hyponyms'] += i.lemma_names()[0] + ', '

                    for j in syn[0].hypernyms():
                        self.result[token]['hyperonyms'] += j.lemma_names()[0] + ', '

            self.show_full_dict()

    def init_main(self):
        self.window = Tk()
        #self.window.geometry("1200x700")
        self.window.title("lab3")

        # self.img = Image.open(r"D:\6sem\Eyazis\3nikita\фон.png")
        # # self.width = 1200
        # # self.height = 700
        # width = 1200
        # height = 700
        # #self.imag = self.img.resize((self.width, self.height), Image.LANCZOS)
        # self.imag = self.img.resize((width, height), Image.LANCZOS)
        # self.image = ImageTk.PhotoImage(self.imag)
        # # self.panel = Label(self.window, image=self.image)
        # # self.panel.pack(side="top", fill="both", expand="no")
        # canvas = Canvas(self.window, width=width, height=height)
        # canvas.pack(side="top", fill="both", expand="no")
        # canvas.create_image(0, 0, anchor="nw", image=self.image)

        self.inputFrame = Frame(self.window, bd=10)
        self.inputText = Text(self.inputFrame, height=10, width=80, wrap=WORD, bd=10)
        self.space = Label(self.inputFrame, text='\n')
        self.createVocabularyButton = Button(self.inputFrame, text='Выполнить анализ', width=25, height=2,
                                             bg='cornflower blue')


        self.createVocabularyButton.config(command=self.viewWindow)

        self.addingFrame = Frame(self.window, bd=40)
        self.addButton = Button(self.addingFrame, text='Текст из файла', width=15, height=2, bg='cornflower blue')
        self.addButton.config(command=self.addTextFromFile)

        self.addingFrame.pack()
        self.addButton.pack(side='left')

        self.addingFrame3 = Frame(self.window, bg='cornflower blue', bd=50)
        self.list_box = Listbox(self.addingFrame3, height=10, width=100)
        self.list_box.grid(row=4, column=0, sticky='nsew', columnspan=3)

        self.addingFrame2 = Frame(self.window, bg='cornflower blue', bd=20)
        self.save_button = Button(self.addingFrame2, text="Сохранить", width=10, height=3, bg='white',
                                  command=self.save_button_click)
        self.upload_button = Button(self.addingFrame2, text="Загрузить", width=10, height=3, bg='white',
                                    command=self.upload_button_click)
        self.update_button = Button(self.addingFrame2, text="Изменить", width=10, height=3, bg='white',
                                    command=self.update_button_click)
        self.search_button = Button(self.addingFrame2, text="Словарь", width=10, height=3, bg='white',
                                    command=self.document_button_click)
        self.back_button = Button(self.addingFrame2, text="^", width=10, height=3, bg='white',
                                  command=self.back_button_click)

        self.inputFrame.pack()
        self.inputText.pack()
        self.space.pack()
        self.createVocabularyButton.pack(side='bottom')
        self.addingFrame3.pack()
        self.addingFrame2.pack()
        self.save_button.pack(side='left')
        self.upload_button.pack(side='left')
        self.update_button.pack(side='left')
        self.search_button.pack(side='left')
        self.back_button.pack(side='left')


def update_empl(string_dict, app):
    string_dict = string_dict.replace('\'', '\"')
    app.curr_result = json.loads(string_dict)
    app.show_curr_dict()

class child_update_empl(Toplevel):
    def __init__(self, root, app, text):
        super().__init__(root)
        self.init_child(text)
        self.view = app

    def init_child(self, text):
        self.title("Изменить")
        self.geometry('800x800+400+300')

        tree_text = Label(self, text='Словарь выглядит следующим образом:')
        tree_text.place(x=50, y=10)

        self.token = Text(self, height=30, width=80, wrap=WORD, bd=20)
        self.token.place(x=50, y=60)
        self.token.insert(1.0, text)

        btn_cancel = Button(self, text='Отмена', command=self.destroy)
        btn_cancel.place(x=600, y=660)

        btn_add = Button(self, text='Изменить')
        btn_add.place(x=490, y=660)
        btn_add.bind('<Button-1>', lambda event: update_empl(self.token.get(1.0, END), self.view))

        self.grab_set()
        self.focus_set()

def document_empl(string_dict):
    string_dict = string_dict.replace('\'', '\"')
    dict = json.loads(string_dict)

    doc = docx.Document()

    for token in dict.keys():
        doc.add_paragraph("Word: " + token)
        if dict[token]['synonyms'] != '':
            doc.add_paragraph("Synonyms: ")
            doc.add_paragraph(dict[token]['synonyms'])
            doc.add_paragraph("\t B этом поле перечислены слова, которые имеют то же или похожее значение, что и слово, которое вы искали.")
        if dict[token]['antonyms'] != '':
            doc.add_paragraph("Antonyms: ")
            doc.add_paragraph(dict[token]['antonyms'])
            doc.add_paragraph("\t B этом поле перечислены слова, имеющие противоположное значение искомому слову.")
        if dict[token]['hyponyms'] != '':
            doc.add_paragraph("Hyponyms: ")
            doc.add_paragraph(dict[token]['hyponyms'])
            doc.add_paragraph("\t B этом поле перечислены слова, которые являются более конкретными, чем слово, которое вы искали. Другими словами, в нем перечислены подтипы или примеры слова.")
        if dict[token]['hyperonyms'] != '':
            doc.add_paragraph("Hyperonyms: ")
            doc.add_paragraph(dict[token]['hyperonyms'])
            doc.add_paragraph("\t B этом поле перечислены слова, которые являются более общими, чем слово, которое вы искали. Другими словами, в нем указана более широкая категория или группа, к которой принадлежит слово.")

    file_path = filedialog.asksaveasfilename(defaultextension=".docx")

    doc.save(file_path)

class child_document_empl(Toplevel):
    def __init__(self, root, app, text):
        super().__init__(root)
        self.init_child(text)
        self.view = app

    def init_child(self, text):
        self.title("Словарь")
        self.geometry('800x800+400+300')

        tree_text = Label(self, text='Словарь выглядит следующим образом:')
        tree_text.place(x=50, y=10)

        self.token = Text(self, height=30, width=80, wrap=WORD, bd=20)
        self.token.place(x=50, y=60)
        self.token.insert(1.0, text)

        btn_cancel = Button(self, text='Отмена', command=self.destroy)
        btn_cancel.place(x=600, y=660)

        btn_add = Button(self, text='Сохранить файл')
        btn_add.place(x=470, y=660)
        btn_add.bind('<Button-1>', lambda event: document_empl(self.token.get(1.0, END)))

        self.grab_set()
        self.focus_set()


if __name__ == '__main__':
    app = Main()
    app.window.mainloop()